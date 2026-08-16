"""
Generate Workforce Management Proposal DOCX for Anna.I
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

# ── Colour constants ──────────────────────────────────────────────
DARK_GREEN  = RGBColor(0x5A, 0x7A, 0x6A)
MID_GREEN   = RGBColor(0x6B, 0x8F, 0x7A)
DARK_SLATE  = RGBColor(0x4A, 0x5E, 0x6A)
BODY_COLOR  = RGBColor(0x33, 0x33, 0x33)
GRAY_TEXT   = RGBColor(0x66, 0x66, 0x66)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_BG = "D9D9D9"
HEADER_GREEN  = "5A7A6A"
ALT_ROW       = "F2F6F4"

OUTPUT_PATH = "/home/z/my-project/docs/staff-roster-proposal.docx"

doc = Document()

# ── Default font ──────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = BODY_COLOR
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Custom Heading styles ─────────────────────────────────────────
for level, (size, color, bold) in {
    1: (18, DARK_GREEN, True),
    2: (14, MID_GREEN, True),
    3: (12, DARK_SLATE, True),
}.items():
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.size = Pt(size)
    hs.font.color.rgb = color
    hs.font.bold = bold
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    hs.paragraph_format.space_after = Pt(8)

# ── Helpers ───────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_GREEN)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = BODY_COLOR
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")  # spacer
    return table


def add_code_block(doc, text, label=None):
    """Add a monospaced code block with light-gray background."""
    if label:
        p_label = doc.add_paragraph()
        r = p_label.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_SLATE
        r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_GRAY_BG}" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = BODY_COLOR
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_body(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    return p


def add_bold_body(doc, bold_text, normal_text=""):
    """Add paragraph with bold lead-in."""
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = BODY_COLOR
    if normal_text:
        r2 = p.add_run(normal_text)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
        r2.font.color.rgb = BODY_COLOR
    return p


# ══════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# Spacer
for _ in range(6):
    doc.add_paragraph("")

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("Staff Roster \u2192 Workforce Management")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = DARK_GREEN
r.font.name = "Calibri"

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Product Strategy Proposal for Anna.I Vendor Portal")
r.font.size = Pt(14)
r.font.color.rgb = GRAY_TEXT
r.font.name = "Calibri"

# Divider line
doc.add_paragraph("")
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_line.add_run("\u2500" * 60)
r.font.color.rgb = MID_GREEN
r.font.size = Pt(10)

# Meta
p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_meta.add_run("Prepared for Product & Engineering Review  |  Confidential")
r.font.size = Pt(11)
r.font.color.rgb = GRAY_TEXT
r.font.name = "Calibri"

# Date
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_date.add_run(datetime.now().strftime("%B %Y"))
r.font.size = Pt(11)
r.font.color.rgb = GRAY_TEXT
r.font.name = "Calibri"

# Page break after cover
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (placeholder)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1.  Executive Summary",
    "2.  Recommended User Journey",
    "3.  Updated Job Workflow",
    "4.  Staff Assignment Workflow",
    "5.  Notification Flow",
    "6.  Customer Experience",
    "7.  Vendor Experience",
    "8.  UI/UX Recommendations",
    "9.  Data Model Considerations",
    "10. Edge Cases and Exception Handling",
    "11. MVP Feature List",
    "12. Future Enhancement Roadmap",
    "13. Final Product Recommendations",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = BODY_COLOR

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

add_body(doc,
    "Anna.I's current Staff Roster module serves as a simple directory \u2014 it stores staff names, "
    "contact details, and role labels for SME vendors. However, the platform's Booking model already "
    "includes an assignedStaffId foreign key, and a backend assign-staff API endpoint exists but remains "
    "entirely unused on the frontend. This disconnect means vendors manage staff in one place while jobs "
    "are dispatched without any link to specific team members."
)

add_body(doc,
    "This proposal outlines the evolution of the Staff Roster into a fully operational Workforce Management "
    "layer that sits naturally between \"Vendor Accepts Job\" and \"Staff Completes Work\" in the task lifecycle. "
    "The MVP focuses on manual staff assignment with customer visibility \u2014 the highest-value, lowest-risk "
    "increment. Subsequent phases introduce staff notifications, skill-based suggestions, auto-assignment, and "
    "performance tracking. This is an enhancement to the existing vendor workflow, not a standalone product."
)

add_body(doc,
    "The business case is clear: customers who see a named, role-identified staff member assigned to their job "
    "experience greater trust and professionalism. Vendors gain accountability, workload visibility, and a "
    "foundation for operational intelligence. The engineering effort is modest because the data model, API, "
    "WebSocket infrastructure, and notification system already exist \u2014 the MVP is primarily about wiring "
    "these pieces together with a well-designed frontend experience."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  2. RECOMMENDED USER JOURNEY
# ══════════════════════════════════════════════════════════════════
doc.add_heading("2. Recommended User Journey", level=1)

add_body(doc, "The updated journey integrates staff assignment at the optimal point in the task lifecycle:")

add_code_block(doc,
    "Customer creates job \u2192 Platform matches vendor \u2192 Vendor notified (15 min timeout)\n"
    "    \u2192 Vendor ACCEPTS job (escrow held) \u2192 [NEW: Vendor assigns staff member]\n"
    "    \u2192 [NEW: Customer sees assigned staff] \u2192 Staff receives assignment notification\n"
    "    \u2192 Vendor/Staff STARTS work \u2192 Work COMPLETED \u2192 Customer VERIFIES \u2192 Escrow RELEASED"
)

doc.add_heading("Why This Placement Is Optimal", level=3)

add_bullet(doc, "Assignment happens AFTER escrow is held (financial commitment confirmed)")
add_bullet(doc, "Assignment happens BEFORE work starts (customer knows who's coming)")
add_bullet(doc, "Vendor has full control over who to assign (no auto-assignment complexity in MVP)")
add_bullet(doc, "Reassignment is possible before work starts without customer disruption")

add_body(doc, "Comparison of current vs. proposed flows:")

add_table(doc,
    headers=["Aspect", "Current Flow", "Proposed Flow"],
    rows=[
        ["Staff awareness", "None", "Notified on assignment"],
        ["Customer visibility", "No staff info", "Sees assigned staff name"],
        ["Vendor control", "Manual name list", "Active assignment to jobs"],
        ["Accountability", "None", "Staff linked to completed bookings"],
    ],
    col_widths=[1.8, 2.4, 2.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  3. UPDATED JOB WORKFLOW
# ══════════════════════════════════════════════════════════════════
doc.add_heading("3. Updated Job Workflow", level=1)

doc.add_heading("SME Vendors (have staff team)", level=2)

add_code_block(doc,
    "1.  Vendor receives TASK_DISPATCHED notification\n"
    "2.  Vendor ACCEPTS \u2192 Booking status: accepted, Escrow: HELD\n"
    "3.  Vendor opens booking \u2192 sees \"Assign Staff\" action\n"
    "4.  Vendor selects staff member from dropdown (filtered to active staff)\n"
    "5.  Booking.assignedStaffId set \u2192 real-time event emitted\n"
    "6.  Customer receives VENDOR_SCHEDULED notification with staff name\n"
    "7.  Vendor/Staff clicks START \u2192 Booking: in_progress, customer: VENDOR_EN_ROUTE\n"
    "8.  Work completed \u2192 verified \u2192 escrow released"
)

doc.add_heading("MICRO Vendors (no staff team)", level=2)

add_body(doc,
    "Same as current flow \u2014 no change. MICRO vendors are the sole operator and do not "
    "see any staff assignment UI elements."
)

doc.add_heading("Key Decision: Staff Assignment Is Optional", level=3)
add_body(doc,
    "Staff assignment is OPTIONAL for SME vendors. They can skip it and proceed as before. "
    "This ensures zero disruption to existing workflows."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  4. STAFF ASSIGNMENT WORKFLOW
# ══════════════════════════════════════════════════════════════════
doc.add_heading("4. Staff Assignment Workflow", level=1)

doc.add_heading("Assignment Dialog (UI)", level=2)

add_bullet(doc, "Triggered from Vendor Schedule page (booking card) and Vendor Task Detail panel")
add_bullet(doc, "Shows a popover/dropdown with active staff members")
add_bullet(doc, "Each staff member shows: name, role badge, current active job count")
add_bullet(doc, '\"Assign & Notify\" button confirms assignment')
add_bullet(doc, "Confirmation toast with staff name")

doc.add_heading("Un-assignment / Reassignment", level=2)

add_bullet(doc, "Available before work starts (booking status: accepted)")
add_bullet(doc, "Change staff \u2192 customer gets updated notification")
add_bullet(doc, "After work starts (in_progress) \u2192 reassignment requires confirmation dialog (shows warning about disruption)")

doc.add_heading("Multiple Staff per Job", level=2)

add_body(doc,
    "NOT recommended for MVP. The current assignedStaffId is a single FK; changing to an array "
    "adds significant complexity. Multiple staff should be a Phase 3 feature requiring a junction "
    "table (StaffAssignment model)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  5. NOTIFICATION FLOW
# ══════════════════════════════════════════════════════════════════
doc.add_heading("5. Notification Flow", level=1)

add_body(doc, "New event-driven notifications leveraging the existing notification infrastructure:")

doc.add_heading("New Event Types", level=2)

add_table(doc,
    headers=["Event", "Trigger", "Recipient", "Channel", "Content"],
    rows=[
        ["STAFF_ASSIGNED", "Vendor assigns staff to booking", "Customer", "WEB_PUSH, WHATSAPP", "Staff name, role, job category"],
        ["STAFF_REASSIGNED", "Vendor changes assigned staff", "Customer", "WEB_PUSH", "New staff name"],
        ["STAFF_UNASSIGNED", "Vendor removes staff", "Customer", "WEB_PUSH", '"Staff assignment updated"'],
    ],
    col_widths=[1.2, 1.5, 0.9, 1.3, 2.1],
)

doc.add_heading("Staff-Side Notifications (Phase 2)", level=2)
add_body(doc,
    "Staff members do not have accounts in MVP, so notifications go to the VENDOR only. "
    "In Phase 2, if the staff contact is a phone number, notifications will be sent via WhatsApp/SMS; "
    "if email, via email."
)

doc.add_heading("Implementation Approach", level=2)
add_bullet(doc, "Leverage existing emitVendorNotification() and WebSocket event system")
add_bullet(doc, 'Add new event type constants in src/lib/constants.ts')
add_bullet(doc, "Customer notification: reuse existing Notification table with recipientType: HOUSEHOLD_MEMBER")
add_bullet(doc, "Real-time: emit via Socket.IO to household room")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  6. CUSTOMER EXPERIENCE
# ══════════════════════════════════════════════════════════════════
doc.add_heading("6. Customer Experience", level=1)

doc.add_heading("Customer SEES", level=2)
add_bullet(doc, 'Staff name (first name + last initial for privacy: e.g. "Ahmad R.")')
add_bullet(doc, 'Staff role badge (e.g. "Technician", "Senior Cleaner")')
add_bullet(doc, '"Your service provider has assigned [Name] to your job" notification')
add_bullet(doc, "Staff name on job detail/tracking page")

doc.add_heading("Customer Does NOT See", level=2)
add_bullet(doc, "Staff full name or personal contact details")
add_bullet(doc, "Staff availability or schedule")
add_bullet(doc, "Internal notes about staff")
add_bullet(doc, "Staff performance metrics")
add_bullet(doc, "Other jobs the staff member is assigned to")
add_bullet(doc, "Whether staff was reassigned (just show current assignee)")

doc.add_heading("Implementation", level=3)
add_body(doc,
    'Add assignedStaff.name to the customer-facing booking detail API response. Only expose name and role \u2014 '
    "never contact, isActive, or other internal fields."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  7. VENDOR EXPERIENCE
# ══════════════════════════════════════════════════════════════════
doc.add_heading("7. Vendor Experience", level=1)

doc.add_heading("Schedule Page Enhancements", level=2)
add_bullet(doc, "Unassigned job indicator: amber badge on booking cards without assigned staff (SME vendors only)")
add_bullet(doc, "Staff avatar/initials on booking cards (when assigned)")
add_bullet(doc, "Quick-assign inline: dropdown directly on booking card without opening detail panel")
add_bullet(doc, 'Filter: "Unassigned" filter tab on schedule')

doc.add_heading("Staff Roster Page Enhancements", level=2)
add_bullet(doc, "Active assignment count per staff member")
add_bullet(doc, '"Currently assigned to: [Job Category] at [Address]" link')
add_bullet(doc, "Workload summary: X active jobs, Y completed today")

doc.add_heading("Task Detail Panel Enhancements", level=2)
add_bullet(doc, '"Assign Staff" section with dropdown + assign button')
add_bullet(doc, "Reassignment flow (for accepted bookings)")
add_bullet(doc, "Un-assign button with confirmation")
add_bullet(doc, "Staff assignment history (audit log of who was assigned/when)")

doc.add_heading("Vendor Dashboard Enhancement (Phase 2)", level=2)
add_bullet(doc, '"Unassigned Jobs" count card')
add_bullet(doc, "Staff workload summary panel")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  8. UI/UX RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("8. UI/UX Recommendations", level=1)

doc.add_heading("8.1 Assign Staff Dialog/Popover", level=2)
add_bold_body(doc, "Location: ", "Booking card (schedule page) and Task Detail panel")
add_bold_body(doc, "Component: ", "Popover (shadcn/ui) \u2014 lightweight, no page navigation")
add_bold_body(doc, "Content: ", "List of active staff with name, role, current workload badge")
add_bold_body(doc, "Interaction: ", "Click staff name \u2192 confirms assignment \u2192 toast")
add_bold_body(doc, "Empty state: ", '"No active staff \u2014 add team members in Staff Roster"')

doc.add_heading("8.2 Staff Roster Screen Enhancements", level=2)
add_bullet(doc, 'Add "Current Jobs" column/badge next to each staff member')
add_bullet(doc, "Link from staff name to filtered schedule view")
add_bullet(doc, "Workload indicator: green (0-1 jobs), amber (2-3), red (4+)")
add_bullet(doc, "Keep existing add/toggle/remove functionality")

doc.add_heading("8.3 Booking Card (Schedule Page)", level=2)
add_bullet(doc, "SME vendors: Show staff avatar circle (initials) when assigned")
add_bullet(doc, 'SME vendors: Show amber "Unassigned" badge when no staff')
add_bullet(doc, "MICRO vendors: No change (they ARE the staff)")

doc.add_heading("8.4 Job Detail Panel", level=2)
add_bullet(doc, 'Add "Assigned Staff" section below booking status')
add_bullet(doc, "Show staff name + role with edit (pencil) icon")
add_bullet(doc, "Edit opens popover to reassign or un-assign")

doc.add_heading("Design Principles", level=2)
add_bullet(doc, "Minimise clicks: assign staff in 1-2 clicks from booking card")
add_bullet(doc, "Reduce mistakes: only show active staff, show current workload")
add_bullet(doc, "Mobile-first: Popover works on mobile, Sheet for task detail")
add_bullet(doc, "Consistent: Use existing Anna.I design system tokens")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  9. DATA MODEL CONSIDERATIONS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("9. Data Model Considerations", level=1)

doc.add_heading("Current Schema (Adequate for MVP)", level=2)
add_bullet(doc, "VendorStaff: Already has name, contact, role, isActive, vendorId")
add_bullet(doc, "Booking.assignedStaffId: Already exists as nullable FK")

add_body(doc, "No schema changes needed for MVP. The existing model supports the full MVP scope.")

doc.add_heading("Phase 2 Schema Additions", level=2)

schema_p2 = """model VendorStaff {
  // Existing fields...
  skills      String?   // JSON array of skill tags
  avatarUrl   String?   // Profile photo
  languages   String?   // JSON array of language codes
  notes       String?   // Internal notes by vendor admin
  branch      String?   // Branch/location identifier

  // New relation for assignment history
  assignmentHistory StaffAssignmentHistory[]
}

model StaffAssignmentHistory {
  id          String   @id @default(cuid())
  bookingId   String
  staffId     String
  vendorId    String
  action      String   // "assigned" | "unassigned" | "reassigned"
  assignedBy  String   // vendorId of who made the change
  notes       String?
  createdAt   DateTime  @default(now())

  booking Booking     @relation(fields: [bookingId], references: [id],
                               onDelete: Cascade)
  staff   VendorStaff @relation(fields: [staffId], references: [id])
  vendor  Vendor      @relation(fields: [vendorId], references: [id])

  @@index([staffId])
  @@index([bookingId])
}"""

add_code_block(doc, schema_p2, label="Prisma Schema \u2014 Phase 2")

doc.add_heading("Phase 3 Schema (Multi-Staff)", level=2)

schema_p3 = """model StaffAssignment {
  id          String   @id @default(cuid())
  bookingId   String
  staffId     String
  role        String   @default("worker") // "lead" | "worker" | "supervisor"
  assignedAt  DateTime @default(now())
  removedAt   DateTime?

  booking Booking     @relation(fields: [bookingId], references: [id])
  staff   VendorStaff @relation(fields: [staffId], references: [id])

  @@index([bookingId])
  @@index([staffId])
}
// This replaces Booking.assignedStaffId with a junction table
// for multiple staff per job."""

add_code_block(doc, schema_p3, label="Prisma Schema \u2014 Phase 3")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  10. EDGE CASES AND EXCEPTION HANDLING
# ══════════════════════════════════════════════════════════════════
doc.add_heading("10. Edge Cases and Exception Handling", level=1)

add_table(doc,
    headers=["Scenario", "Handling", "User Impact"],
    rows=[
        ["Staff deleted while assigned to active booking",
         'Prevent deletion if staff has active bookings. Show error: "Cannot remove staff with X active assignments"',
         "Vendor must reassign before removing"],
        ["Staff deactivated (isActive=false) while assigned",
         "Keep existing assignment. Mark with visual indicator. Vendor prompted to reassign",
         "Customer still sees staff, but vendor warned"],
        ["Vendor rejects job after assigning staff",
         "Assignment cleared automatically on booking cancel. Staff \"unassigned\" event fired",
         "Clean state, no orphan assignments"],
        ["Job times out (no vendor response)",
         "N/A \u2014 assignment only happens after vendor accepts",
         "No impact"],
        ["Staff has 5+ concurrent jobs",
         "Show amber/red workload badge. Optional: soft warning on assign",
         "Vendor makes informed decision"],
        ["MICRO vendor tries to access staff features",
         "UI elements hidden via vendorType check. No staff tab shown",
         "Zero disruption"],
        ["Reassignment after work started (in_progress)",
         "Show confirmation dialog. Log reassignment in audit. Notify customer",
         "Allows flexibility with accountability"],
        ["Network failure during assignment",
         "Optimistic UI with rollback. React Query retry. Local state reset on error",
         "Graceful degradation"],
        ["Customer disputes job with assigned staff",
         "Dispute flow unchanged. Staff info preserved for ops review",
         "No change to dispute process"],
        ["Vendor account suspended",
         "All staff assignments preserved but frozen. No new assignments",
         "Data integrity maintained"],
    ],
    col_widths=[2.0, 2.7, 2.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  11. MVP FEATURE LIST
# ══════════════════════════════════════════════════════════════════
doc.add_heading("11. MVP Feature List", level=1)

doc.add_heading("MVP Scope (Immediate Value)", level=2)

mvp_items = [
    ("\u2705 Assign Staff UI: ", "Popover on booking cards and task detail (SME only)"),
    ("\u2705 Un-assign Staff: ", "Clear assignedStaffId with confirmation"),
    ("\u2705 Reassign Staff: ", "Change assigned staff before work starts"),
    ("\u2705 Customer Visibility: ", "Show assigned staff name on booking details"),
    ("\u2705 Customer Notification: ", "STAFF_ASSIGNED event via WebSocket + Notification table"),
    ("\u2705 Unassigned Job Indicator: ", "Amber badge on schedule for unassigned SME bookings"),
    ("\u2705 Staff Workload Display: ", "Active job count on staff roster"),
    ("\u2705 Staff Delete Guard: ", "Prevent removal if staff has active bookings"),
]
for bold_part, normal_part in mvp_items:
    add_bold_body(doc, bold_part, normal_part)

doc.add_heading("NOT in MVP", level=2)

non_mvp = [
    "Staff accept/decline flow (staff have no accounts)",
    "Staff notifications (WhatsApp/SMS to staff)",
    "Skill-based assignment suggestions",
    "Auto-assignment algorithms",
    "Staff performance dashboards",
    "Multiple staff per job",
    "Staff profiles with photos/skills/certifications",
    "Staff mobile app",
]
for item in non_mvp:
    add_bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  12. FUTURE ENHANCEMENT ROADMAP
# ══════════════════════════════════════════════════════════════════
doc.add_heading("12. Future Enhancement Roadmap", level=1)

doc.add_heading("Phase 2 \u2014 Operational Intelligence (After MVP Validates)", level=2)

phase2 = [
    "Staff notification channel (WhatsApp/SMS/email based on contact type)",
    "Staff accept/decline response flow",
    "Skills and certifications on staff profiles",
    'Skill-based assignment suggestions (“Recommended: Ahmad (AC specialist)”)',
    "Workload-balanced suggestions",
    "Assignment history and audit trail",
    "Vendor dashboard: unassigned jobs card, staff workload summary",
    "Staff availability scheduling (shifts, leave, OOO)",
    "Zone-based staff matching (staff who serve the customer's area)",
]
for item in phase2:
    add_bullet(doc, item)

doc.add_heading("Phase 3 \u2014 Automation & Scale (When Vendor Base Grows)", level=2)

phase3 = [
    "Auto-assignment rules engine (configurable by vendor)",
    "Multi-staff job support (lead + worker roles)",
    "Staff performance metrics (completion rate, avg rating, response time)",
    "Staff leaderboard and ranking within vendor organisation",
    "Predictive staffing suggestions (ML-based workload forecasting)",
    "Staff mobile app (native or PWA) for field updates",
    "Customer preferred staff selection",
    "Staff scheduling and shift management",
]
for item in phase3:
    add_bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  13. FINAL PRODUCT RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("13. Final Product Recommendations", level=1)

doc.add_heading("Strategic Recommendations", level=2)

recs = [
    ("1. Start with SME vendors only. ",
     'MICRO vendors don\'t need staff features. Gate all staff UI behind vendorType === "SME" check \u2014 '
     "already pattern established in codebase."),
    ("2. Manual assignment first. ",
     "Auto-assignment sounds appealing but adds complexity (routing at staff level, availability conflicts, "
     "workload balancing). Let vendors build the habit of assigning staff manually first."),
    ("3. Leverage existing infrastructure. ",
     "The codebase already has: WebSocket real-time events, Notification table, assignedStaffId FK, "
     "assign-staff API, read-only staff display. MVP is mostly wiring up existing pieces."),
    ("4. Customer visibility is the highest-value MVP feature. ",
     'Customers currently see no staff information. Showing "Ahmad R. (Technician)" on their job adds '
     "trust and professionalism with minimal effort."),
    ("5. Design for data, not just UI. ",
     "Every assignment/reassignment creates an audit trail. This data becomes the foundation for "
     "Phase 2/3 intelligence features."),
    ("6. Zero disruption to existing flows. ",
     "Staff assignment is OPTIONAL. Vendors who don't assign staff experience no change. MICRO vendors "
     "see no difference."),
]
for bold_part, normal_part in recs:
    add_bold_body(doc, bold_part, normal_part)

add_body(doc, "")  # spacer

doc.add_heading("Implementation Priority", level=2)

priority = [
    "Wire up assign-staff UI to existing API (1-2 days)",
    "Add un-assign endpoint + guard on staff delete (0.5 day)",
    "Customer notification event + WebSocket emit (1 day)",
    "Customer-facing staff display (1 day)",
    "Schedule page unassigned indicator + workload (1 day)",
    "Edge case handling + testing (1 day)",
]
for i, item in enumerate(priority, 1):
    add_bullet(doc, f"{i}. {item}")

add_body(doc, "")
p_est = doc.add_paragraph()
r = p_est.add_run("Total MVP Estimate: ~5-6 engineering days")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = DARK_GREEN
r.font.name = "Calibri"

add_body(doc, "")
doc.add_heading("Risk Assessment", level=2)

add_table(doc,
    headers=["Risk", "Likelihood", "Impact", "Mitigation"],
    rows=[
        ["Low vendor adoption", "Medium", "High",
         "Make assignment optional; show value through customer trust signals"],
        ["Staff contact data quality", "High", "Medium",
         "Validate contact format; allow vendor to update anytime"],
        ["Reassignment confusion", "Low", "Medium",
         "Clear confirmation dialogs; audit trail"],
        ["Scope creep", "Medium", "High",
         "Strict MVP boundary; no staff accounts, no auto-assign"],
    ],
    col_widths=[1.6, 1.0, 0.9, 3.5],
)

# ── Add page numbers via footer ──────────────────────────────────
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Field code for page number
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    # Confidential notice
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Anna.I \u2014 Confidential")
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY_TEXT
    r.font.name = "Calibri"

# ── Save ──────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
